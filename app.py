import io
import logging
import os
import time
import zipfile
from datetime import datetime
from typing import Any

import streamlit as st

try:
    from openai import OpenAI
except ImportError:
    OpenAI = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx
except ImportError:
    docx = None

APP_TITLE = "AI 简历优化与面试模拟 Agent"
LOGGER = logging.getLogger(__name__)
ALLOWED_MODELS = ("gpt-5.2", "gpt-5.2-mini", "gpt-5.2-nano")
DEFAULT_MODEL = os.getenv("OPENAI_MODEL", "gpt-5.2")
if DEFAULT_MODEL not in ALLOWED_MODELS:
    DEFAULT_MODEL = ALLOWED_MODELS[0]
RESULT_KEYS = (
    "analysis_result",
    "interview_questions",
    "interview_feedback",
    "final_report",
)
MAX_UPLOAD_BYTES = 5 * 1024 * 1024
MAX_PDF_PAGES = 20
MAX_DOCX_PARAGRAPHS = 500
MAX_DOCX_FILES = 200
MAX_DOCX_ENTRY_BYTES = 5 * 1024 * 1024
MAX_DOCX_UNCOMPRESSED_BYTES = 15 * 1024 * 1024
MAX_DOCX_COMPRESSION_RATIO = 100
MAX_EXTRACTED_CHARS = 30_000
MAX_RESUME_CHARS = 30_000
MAX_JD_CHARS = 20_000
MAX_QUESTION_CHARS = 5_000
MAX_ANSWER_CHARS = 10_000
MAX_AI_CALLS_PER_SESSION = 12
SESSION_DEFAULTS: dict[str, Any] = {
    "logs": [],
    "ai_call_count": 0,
    "resume_text": "",
    "jd_text": "",
    "analysis_result": "",
    "interview_questions": "",
    "interview_feedback": "",
    "final_report": "",
}


class AiServiceError(RuntimeError):
    pass


SYSTEM_PROMPT = """
你是一个专业的 AI 简历优化与面试模拟 Agent。
你需要帮助用户完成：简历诊断、岗位匹配、简历优化、面试问题生成、面试回答评分。

硬性规则：
1. 只能基于用户提供的简历和岗位 JD 分析，不得编造经历。
2. 可以优化表达，但不能虚构公司、学校、项目、证书、奖项、数字成果。
3. 如果缺少量化指标，只能写“建议补充：xxx 指标”，不能直接编造数字。
4. 简历优化要遵循 STAR 法则：Situation、Task、Action、Result。
5. 输出要结构化，适合直接展示在网页 Demo 中。
6. 语言专业、直接、具体，不要空泛鼓励。
"""

ANALYSIS_PROMPT_TEMPLATE = """
请根据以下简历和目标岗位 JD，完成一次完整的简历优化分析。

【简历内容】
{resume}

【目标岗位 JD】
{jd}

请按以下结构输出：

# AI 简历诊断与岗位匹配报告

## 1. 总体匹配度
- 匹配度评分：0-100
- 结论：高匹配 / 中等匹配 / 低匹配
- 一句话原因

## 2. 分项评分
- 技能匹配度：0-100
- 项目经验匹配度：0-100
- 关键词覆盖度：0-100
- 表达质量：0-100
- 面试风险：0-100，分数越高代表风险越低

## 3. 岗位 JD 关键词
提取 10-20 个应该出现在简历中的关键词。

## 4. 简历优势
列出 3-5 条。

## 5. 简历主要问题
列出 3-5 条，必须具体。

## 6. 逐段优化建议
针对个人摘要、技能模块、工作经历、项目经历分别给出修改建议。

## 7. 可直接替换的简历表达
给出 3-5 条优化后的简历 bullet。
注意：不得编造用户没有提供的经历或数字。没有数字时写“建议补充：xxx 指标”。

## 8. 投递前检查清单
给出 5 条检查项。
"""

INTERVIEW_PROMPT_TEMPLATE = """
请根据以下简历和目标岗位 JD，生成一轮定制化模拟面试题。

【简历内容】
{resume}

【目标岗位 JD】
{jd}

请按以下结构输出：

# 模拟面试题

## 1. 自我介绍类
生成 2 题。

## 2. 项目深挖类
生成 5 题，必须结合简历中的具体经历。

## 3. 技术/专业能力类
生成 5 题，必须结合 JD 要求。

## 4. 岗位匹配类
生成 3 题。

## 5. 行为面试类
生成 3 题。

## 6. 压力面试类
生成 2 题。

每道题都要包含：
- 问题
- 考察点
- 回答建议
"""

SCORING_PROMPT_TEMPLATE = """
请对用户的面试回答进行评分和反馈。

【目标岗位 JD】
{jd}

【简历内容】
{resume}

【面试问题】
{question}

【用户回答】
{answer}

请按以下结构输出：

# 面试回答评分

## 1. 总分
给出 0-100 分。

## 2. 分项评分
- 内容完整度：0-20
- 逻辑结构：0-20
- 岗位相关性：0-20
- 项目细节可信度：0-20
- 表达清晰度：0-20

## 3. 主要问题
指出 2-4 个具体问题。

## 4. 优化建议
给出可执行建议。

## 5. 参考回答
基于用户已有简历和原始回答，生成一个更好的回答版本。
不得编造经历或数字。缺少数字时写“建议补充：xxx 指标”。

## 6. 下一轮追问
生成 1-2 个面试官可能继续追问的问题。
"""

REPORT_PROMPT_TEMPLATE = """
请根据以下信息生成一份完整的求职优化报告。

【简历内容】
{resume}

【目标岗位 JD】
{jd}

【简历分析结果】
{analysis}

【面试题】
{questions}

【面试反馈】
{feedback}

请输出 Markdown 报告，结构如下：

# AI 求职优化报告

## 1. 岗位目标
## 2. 简历匹配度总结
## 3. 简历最需要修改的问题
## 4. 可直接执行的简历优化建议
## 5. 模拟面试表现分析
## 6. 高频风险问题
## 7. 3 天行动计划
## 8. 7 天行动计划

要求：
1. 内容具体。
2. 不虚构用户经历。
3. 适合直接导出提交或保存。
"""

SAMPLE_RESUME = """
张三｜AI 产品经理｜3 年经验

教育背景：
XX 大学，计算机科学与技术，本科。

工作经历：
某科技公司 AI 产品经理，2022.07-至今
- 负责 AI 助手产品的需求分析、原型设计和上线跟进。
- 与算法、后端、前端团队协作，推动知识库问答、智能推荐等功能落地。
- 参与用户调研和数据分析，持续优化产品体验。

项目经历：
企业知识库问答系统
- 负责需求调研、产品方案设计和验收测试。
- 协调算法团队接入向量检索和大模型问答能力。
- 输出 PRD、流程图和测试用例。
- 建议补充：用户使用量、回答准确率、人工成本节省等指标。

技能：
AI 产品设计、用户调研、PRD、Axure、SQL、数据分析、大模型应用、RAG。
""".strip()

SAMPLE_JD = """
岗位：AI 产品经理

岗位职责：
1. 负责 AI Agent、知识库问答、智能助手等产品的规划和落地。
2. 分析用户需求，设计产品方案、业务流程和交互原型。
3. 协调算法、工程、设计和运营团队推动项目上线。
4. 跟踪产品数据，持续优化用户体验和业务效果。

任职要求：
1. 2 年以上 AI 产品或 ToB 产品经验。
2. 熟悉大模型、RAG、Agent、Prompt Engineering 等技术概念。
3. 具备较强的数据分析、需求抽象和项目推进能力。
4. 能够独立输出 PRD、流程图、原型和项目复盘。
5. 有企业知识库、智能客服、办公自动化产品经验优先。
""".strip()


def configure_page() -> None:
    st.set_page_config(
        page_title=APP_TITLE,
        page_icon="🧑‍💼",
        layout="wide",
    )


def init_state() -> None:
    for key, value in SESSION_DEFAULTS.items():
        if key not in st.session_state:
            st.session_state[key] = value.copy() if isinstance(value, list) else value


def add_log(message: str) -> None:
    timestamp = datetime.now().strftime("%H:%M:%S")
    st.session_state.logs.append(f"[{timestamp}] {message}")


def enforce_text_limit(label: str, text: str, max_chars: int) -> None:
    if len(text) > max_chars:
        raise ValueError(f"{label}超过 {max_chars} 字符，请压缩内容后再提交。")


def validate_ai_inputs(fields: list[tuple[str, str, int]]) -> bool:
    for label, text, max_chars in fields:
        try:
            enforce_text_limit(label, text, max_chars)
        except ValueError as exc:
            st.error(str(exc))
            return False
    return True


def enforce_ai_quota() -> bool:
    if st.session_state.ai_call_count >= MAX_AI_CALLS_PER_SESSION:
        st.error(f"本会话已达到 {MAX_AI_CALLS_PER_SESSION} 次 AI 调用上限，请稍后刷新会话或使用演示模式。")
        return False
    return True


def record_ai_call() -> None:
    st.session_state.ai_call_count += 1


def limit_extracted_text(text: str) -> str:
    enforce_text_limit("上传文件提取内容", text, MAX_EXTRACTED_CHARS)
    return text


def validate_model(model: str) -> str:
    if model not in ALLOWED_MODELS:
        raise ValueError("模型不在允许列表中，请在侧边栏重新选择。")
    return model


def validate_docx_zip(raw: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as archive:
            infos = archive.infolist()
    except zipfile.BadZipFile as exc:
        raise ValueError("DOCX 文件格式无效，请更换文件或直接粘贴简历文本。") from exc

    if len(infos) > MAX_DOCX_FILES:
        raise ValueError("DOCX 文件结构过大，请精简后再上传。")

    total_size = 0
    for info in infos:
        total_size += info.file_size
        if info.file_size > MAX_DOCX_ENTRY_BYTES:
            raise ValueError("DOCX 文件包含过大的内部条目，请精简后再上传。")
        if total_size > MAX_DOCX_UNCOMPRESSED_BYTES:
            raise ValueError("DOCX 解压后内容过大，请精简后再上传。")
        if info.compress_size and info.file_size / info.compress_size > MAX_DOCX_COMPRESSION_RATIO:
            raise ValueError("DOCX 文件压缩比例异常，请更换文件或直接粘贴简历文本。")


def read_uploaded_file(uploaded_file: Any) -> str:
    if uploaded_file is None:
        return ""

    filename = uploaded_file.name.lower()
    file_size = getattr(uploaded_file, "size", None)
    if file_size is not None and file_size > MAX_UPLOAD_BYTES:
        raise ValueError("上传文件不能超过 5 MB，请压缩文件或直接粘贴简历文本。")

    uploaded_file.seek(0)
    raw = uploaded_file.read(MAX_UPLOAD_BYTES + 1)
    if len(raw) > MAX_UPLOAD_BYTES:
        raise ValueError("上传文件不能超过 5 MB，请压缩文件或直接粘贴简历文本。")

    if filename.endswith((".txt", ".md")):
        return limit_extracted_text(raw.decode("utf-8", errors="ignore"))

    if filename.endswith(".pdf"):
        if PdfReader is None:
            raise ValueError("PDF 解析库未安装。请安装 pypdf，或直接粘贴简历文本。")
        try:
            reader = PdfReader(io.BytesIO(raw))
            if reader.is_encrypted:
                raise ValueError("PDF 文件已加密，暂时无法解析。请解除加密后上传，或直接粘贴简历文本。")
            if len(reader.pages) > MAX_PDF_PAGES:
                raise ValueError(f"PDF 页数不能超过 {MAX_PDF_PAGES} 页，请精简后再上传。")
            return limit_extracted_text("\n".join(page.extract_text() or "" for page in reader.pages))
        except ValueError:
            raise
        except Exception as exc:
            LOGGER.exception("PDF 解析失败")
            raise ValueError("PDF 解析失败，请检查文件是否损坏，或直接粘贴简历文本。") from exc

    if filename.endswith(".docx"):
        if docx is None:
            raise ValueError("DOCX 解析库未安装。请安装 python-docx，或直接粘贴简历文本。")
        try:
            validate_docx_zip(raw)
            document = docx.Document(io.BytesIO(raw))
            if len(document.paragraphs) > MAX_DOCX_PARAGRAPHS:
                raise ValueError(f"DOCX 段落数不能超过 {MAX_DOCX_PARAGRAPHS} 段，请精简后再上传。")
            return limit_extracted_text("\n".join(paragraph.text for paragraph in document.paragraphs))
        except ValueError:
            raise
        except Exception as exc:
            LOGGER.exception("DOCX 解析失败")
            raise ValueError("DOCX 解析失败，请检查文件格式是否正确，或直接粘贴简历文本。") from exc

    return limit_extracted_text(raw.decode("utf-8", errors="ignore"))


def resolve_api_key(api_key: str) -> str:
    return api_key.strip()


def get_client(api_key: str) -> Any:
    resolved_api_key = resolve_api_key(api_key)
    if OpenAI is None:
        raise RuntimeError("openai 库未安装，请先安装 openai。")
    if not resolved_api_key:
        raise RuntimeError("缺少 OpenAI API Key。请在侧边栏输入，或开启演示模式。")
    return OpenAI(api_key=resolved_api_key)


def extract_response_text(response: Any) -> str:
    status = getattr(response, "status", None)
    if status and status != "completed":
        raise ValueError("AI 响应未完成。")

    output_text = str(getattr(response, "output_text", "")).strip()
    if output_text:
        return output_text

    try:
        fallback_text = str(response.output[0].content[0].text).strip()
    except (AttributeError, IndexError, TypeError) as exc:
        raise ValueError("AI 响应中没有可用文本。") from exc

    if not fallback_text:
        raise ValueError("AI 响应中没有可用文本。")
    return fallback_text


def call_ai(
    api_key: str,
    model: str,
    system_prompt: str,
    user_prompt: str,
    max_output_tokens: int = 3000,
) -> str:
    client = get_client(api_key)
    safe_model = validate_model(model)
    response = client.responses.create(
        model=safe_model,
        input=[
            {"role": "developer", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        max_output_tokens=max_output_tokens,
    )
    return extract_response_text(response)


def safe_call_ai(
    api_key: str,
    model: str,
    user_prompt: str,
    max_output_tokens: int,
) -> str:
    try:
        return call_ai(
            api_key=api_key,
            model=model,
            system_prompt=SYSTEM_PROMPT,
            user_prompt=user_prompt,
            max_output_tokens=max_output_tokens,
        )
    except Exception as exc:
        LOGGER.exception("AI 调用失败")
        raise AiServiceError("AI 服务暂时不可用，请稍后重试，或检查 API Key 与模型配置。") from exc


def run_ai_request(
    api_key: str,
    model: str,
    user_prompt: str,
    max_output_tokens: int,
) -> str | None:
    try:
        result = safe_call_ai(
            api_key=api_key,
            model=model,
            user_prompt=user_prompt,
            max_output_tokens=max_output_tokens,
        )
    except AiServiceError as exc:
        st.error(str(exc))
        return None
    record_ai_call()
    return result


def demo_analysis() -> str:
    return """
# 简历与岗位匹配分析 Demo

## 1. 总体匹配度

- 匹配度评分：78/100
- 结论：中等偏高匹配
- 主要原因：简历中已有相关项目和技能基础，但岗位关键词覆盖不足，项目成果量化不够。

## 2. 匹配优势

1. 简历中已有与目标岗位相关的项目经验。
2. 技能栈与 JD 中部分核心要求存在重合。
3. 项目经历具备可包装为 STAR 结构的基础。
4. 如果补充量化指标，简历竞争力会明显提升。

## 3. 主要缺口

1. 简历顶部缺少针对目标岗位的个人摘要。
2. 项目经历更多描述“做了什么”，缺少“产生什么结果”。
3. JD 中的关键词没有系统嵌入到技能和项目模块。
4. 缺少可验证的业务指标或性能指标。

## 4. 简历优化建议

### 个人摘要建议

建议加入 80-120 字个人摘要，突出目标岗位、核心技能、项目经验和可量化成果。

### 项目经历优化方向

原始表达通常容易写成：

> 负责系统开发，完成相关功能。

建议改为：

> 基于 XX 技术实现 XX 模块，解决 XX 问题，支持 XX 场景；建议补充性能提升、效率提升、用户规模或业务转化等量化指标。

### 技能模块建议

将技能分为：

- 编程语言 / 工具
- 框架 / 平台
- 数据处理 / 分析能力
- 项目管理 / 协作能力

## 5. 下一步建议

1. 先补充 2-3 个项目的量化结果。
2. 将 JD 高频关键词嵌入简历摘要、技能模块和项目描述。
3. 为每个项目准备 3 个面试追问答案。
"""


def demo_questions() -> str:
    return """
# 定制化模拟面试题 Demo

## 自我介绍类

1. 请用 1 分钟介绍你自己，并说明你为什么适合这个岗位。
2. 你简历中最能证明岗位匹配度的经历是哪一段？

## 项目深挖类

1. 你在核心项目中具体负责哪一部分？
2. 这个项目最难的问题是什么？你如何解决？
3. 项目中有没有可以量化的结果？
4. 如果重新做一次，你会如何优化？
5. 你如何证明这个项目是你真实参与的？

## 岗位匹配类

1. JD 中提到的核心能力，你最匹配哪几项？
2. 你目前和岗位要求最大的差距是什么？
3. 入职后 30 天，你会如何快速上手？
"""


def demo_feedback() -> str:
    return """
# 面试回答评分 Demo

## 1. 总分

72/100

## 2. 分项评分

- 内容完整度：15/20
- 逻辑结构：14/20
- 岗位相关性：16/20
- 项目细节可信度：13/20
- 表达清晰度：14/20

## 3. 主要问题

1. 回答中有项目背景，但缺少明确的任务目标。
2. 对个人贡献描述不够具体。
3. 没有补充结果指标，例如效率提升、成本降低、准确率提升或用户反馈。
4. 回答结构可以更接近 STAR 法则。

## 4. 优化建议

建议按照以下结构回答：

1. 项目背景：为什么做这个项目。
2. 你的任务：你负责什么。
3. 具体行动：你用了什么方法。
4. 结果证明：项目产生了什么效果。
5. 岗位关联：这个经历如何证明你适合目标岗位。

## 5. 参考回答模板

我参与的这个项目主要是为了解决 XX 问题。当时我的核心任务是 XX。我主要做了三件事：第一，XX；第二，XX；第三，XX。最终项目实现了 XX 效果。虽然目前简历中还没有写出完整量化指标，但我建议补充例如处理效率、准确率、用户规模或业务节省时间等数据。这个经历能够体现我在目标岗位中需要的 XX 能力。
"""


def demo_report() -> str:
    return f"""
# AI 求职优化报告 Demo

## 1. 岗位目标

目标岗位基于用户输入 JD，系统已完成简历匹配和面试训练。

## 2. 简历匹配度总结

{st.session_state.analysis_result or "尚未生成简历分析结果。"}

## 3. 模拟面试题

{st.session_state.interview_questions or "尚未生成面试题。"}

## 4. 面试反馈

{st.session_state.interview_feedback or "尚未进行面试回答评分。"}

## 5. 下一步行动计划

1. 补充项目中的量化指标。
2. 将 JD 关键词嵌入简历摘要、技能和项目经历。
3. 准备 3 个项目深挖问题答案。
4. 每天完成 2-3 道模拟面试题训练。
"""


def render_sidebar() -> tuple[str, str, bool]:
    with st.sidebar:
        st.header("配置")
        api_key = st.text_input(
            "OpenAI API Key",
            value="",
            type="password",
            help="请输入个人 API Key；服务端不会把环境变量或 Secrets 中的 Key 下发给用户。",
        )
        model = st.selectbox(
            "模型名称",
            options=ALLOWED_MODELS,
            index=ALLOWED_MODELS.index(DEFAULT_MODEL),
            help="仅允许使用白名单模型，避免误用高成本或未知模型。",
        )
        demo_mode = st.toggle("演示模式：无 API Key 也可跑通页面", value=False)

        st.divider()
        if st.button("清空结果"):
            st.session_state.logs = []
            for key in RESULT_KEYS:
                st.session_state[key] = ""
            st.success("已清空。")

        st.divider()
        st.markdown("### Demo 说明")
        st.markdown(
            """
            适合活动录屏展示：
            1. 上传/粘贴简历
            2. 输入岗位 JD
            3. 一键分析
            4. 生成面试题
            5. 输入回答并评分
            6. 导出报告
            """
        )
    return api_key, model, demo_mode


def render_input_tab() -> None:
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("简历输入")
        uploaded_file = st.file_uploader("上传简历文件", type=["pdf", "docx", "txt", "md"])
        if uploaded_file is not None:
            try:
                extracted = read_uploaded_file(uploaded_file)
            except ValueError as exc:
                st.error(str(exc))
                add_log(f"[ResumeParserAgent] 文件读取失败：{uploaded_file.name}")
            else:
                if extracted:
                    st.session_state.resume_text = extracted
                    add_log(f"[ResumeParserAgent] 已读取文件：{uploaded_file.name}")

        st.session_state.resume_text = st.text_area(
            "或直接粘贴简历文本",
            value=st.session_state.resume_text,
            height=360,
            placeholder="粘贴你的简历文本...",
        )

    with col2:
        st.subheader("目标岗位 JD")
        st.session_state.jd_text = st.text_area(
            "粘贴目标岗位 JD",
            value=st.session_state.jd_text,
            height=360,
            placeholder="粘贴招聘 JD，包括岗位职责、任职要求、加分项...",
        )
        if st.button("填入示例数据"):
            st.session_state.resume_text = SAMPLE_RESUME
            st.session_state.jd_text = SAMPLE_JD
            add_log("[Demo] 已填入示例简历和 JD。")
            st.rerun()


def render_analysis_tab(api_key: str, model: str, demo_mode: bool) -> None:
    st.subheader("简历诊断与岗位匹配")
    if st.button("开始分析简历", type="primary"):
        resume = st.session_state.resume_text.strip()
        jd = st.session_state.jd_text.strip()
        if not resume or not jd:
            st.error("请先输入简历和目标岗位 JD。")
        elif not validate_ai_inputs([
            ("简历内容", resume, MAX_RESUME_CHARS),
            ("目标岗位 JD", jd, MAX_JD_CHARS),
        ]):
            return
        elif not demo_mode and not enforce_ai_quota():
            return
        else:
            add_log("[ResumeParserAgent] 正在解析简历...")
            add_log("[JDAnalyzerAgent] 正在分析岗位 JD...")
            add_log("[MatchAgent] 正在计算岗位匹配度...")
            add_log("[ResumeOptimizerAgent] 正在生成简历优化建议...")
            with st.spinner("Agent 正在分析简历与 JD..."):
                time.sleep(1 if demo_mode else 0)
                result = (
                    demo_analysis()
                    if demo_mode
                    else run_ai_request(
                        api_key=api_key,
                        model=model,
                        user_prompt=ANALYSIS_PROMPT_TEMPLATE.format(resume=resume, jd=jd),
                        max_output_tokens=4000,
                    )
                )
            if result is None:
                return
            st.session_state.analysis_result = result
            add_log("[Done] 简历分析完成。")

    if st.session_state.analysis_result:
        st.markdown(st.session_state.analysis_result)
        st.download_button(
            "下载简历分析结果 Markdown",
            data=st.session_state.analysis_result,
            file_name="resume_analysis.md",
            mime="text/markdown",
        )
    else:
        st.info("点击“开始分析简历”后，这里会显示匹配度、问题诊断和优化建议。")


def render_interview_tab(api_key: str, model: str, demo_mode: bool) -> None:
    st.subheader("模拟面试")
    col_a, col_b = st.columns([1, 1])

    with col_a:
        if st.button("生成模拟面试题"):
            resume = st.session_state.resume_text.strip()
            jd = st.session_state.jd_text.strip()
            if not resume or not jd:
                st.error("请先输入简历和目标岗位 JD。")
            elif not validate_ai_inputs([
                ("简历内容", resume, MAX_RESUME_CHARS),
                ("目标岗位 JD", jd, MAX_JD_CHARS),
            ]):
                return
            elif not demo_mode and not enforce_ai_quota():
                return
            else:
                add_log("[InterviewAgent] 正在生成定制化面试题...")
                with st.spinner("正在生成面试题..."):
                    time.sleep(1 if demo_mode else 0)
                    questions = (
                        demo_questions()
                        if demo_mode
                        else run_ai_request(
                            api_key=api_key,
                            model=model,
                            user_prompt=INTERVIEW_PROMPT_TEMPLATE.format(resume=resume, jd=jd),
                            max_output_tokens=3500,
                        )
                    )
                if questions is None:
                    return
                st.session_state.interview_questions = questions
                add_log("[Done] 面试题生成完成。")

        if st.session_state.interview_questions:
            st.markdown(st.session_state.interview_questions)
            st.download_button(
                "下载面试题 Markdown",
                data=st.session_state.interview_questions,
                file_name="interview_questions.md",
                mime="text/markdown",
            )
        else:
            st.info("点击“生成模拟面试题”后，这里会展示问题。")

    with col_b:
        st.markdown("### 回答评分")
        question = st.text_area(
            "选择或粘贴一道面试问题",
            height=120,
            placeholder="例如：请介绍一个你最能体现岗位匹配度的项目。",
        )
        answer = st.text_area("输入你的回答", height=220, placeholder="输入你的面试回答...")

        if st.button("评分并给出改进建议"):
            resume = st.session_state.resume_text.strip()
            jd = st.session_state.jd_text.strip()
            if not question.strip() or not answer.strip():
                st.error("请先输入面试问题和你的回答。")
            elif not resume or not jd:
                st.error("请先输入简历和目标岗位 JD。")
            elif not validate_ai_inputs([
                ("简历内容", resume, MAX_RESUME_CHARS),
                ("目标岗位 JD", jd, MAX_JD_CHARS),
                ("面试问题", question, MAX_QUESTION_CHARS),
                ("面试回答", answer, MAX_ANSWER_CHARS),
            ]):
                return
            elif not demo_mode and not enforce_ai_quota():
                return
            else:
                add_log("[ScoringAgent] 正在评估面试回答...")
                with st.spinner("正在评分..."):
                    time.sleep(1 if demo_mode else 0)
                    feedback = (
                        demo_feedback()
                        if demo_mode
                        else run_ai_request(
                            api_key=api_key,
                            model=model,
                            user_prompt=SCORING_PROMPT_TEMPLATE.format(
                                resume=resume,
                                jd=jd,
                                question=question,
                                answer=answer,
                            ),
                            max_output_tokens=3000,
                        )
                    )
                if feedback is None:
                    return
                st.session_state.interview_feedback = feedback
                add_log("[Done] 面试回答评分完成。")

        if st.session_state.interview_feedback:
            st.markdown(st.session_state.interview_feedback)


def render_report_tab(api_key: str, model: str, demo_mode: bool) -> None:
    st.subheader("报告与运行日志")
    if st.button("生成完整求职优化报告", type="primary"):
        resume = st.session_state.resume_text.strip()
        jd = st.session_state.jd_text.strip()
        if not resume or not jd:
            st.error("请先输入简历和目标岗位 JD。")
        elif not validate_ai_inputs([
            ("简历内容", resume, MAX_RESUME_CHARS),
            ("目标岗位 JD", jd, MAX_JD_CHARS),
            ("简历分析结果", st.session_state.analysis_result, MAX_EXTRACTED_CHARS),
            ("面试题", st.session_state.interview_questions, MAX_EXTRACTED_CHARS),
            ("面试反馈", st.session_state.interview_feedback, MAX_EXTRACTED_CHARS),
        ]):
            return
        elif not demo_mode and not enforce_ai_quota():
            return
        else:
            add_log("[ReportAgent] 正在生成求职优化报告...")
            with st.spinner("正在生成报告..."):
                time.sleep(1 if demo_mode else 0)
                report = (
                    demo_report()
                    if demo_mode
                    else run_ai_request(
                        api_key=api_key,
                        model=model,
                        user_prompt=REPORT_PROMPT_TEMPLATE.format(
                            resume=resume,
                            jd=jd,
                            analysis=st.session_state.analysis_result,
                            questions=st.session_state.interview_questions,
                            feedback=st.session_state.interview_feedback,
                        ),
                        max_output_tokens=5000,
                    )
                )
            if report is None:
                return
            st.session_state.final_report = report
            add_log("[Done] 报告生成完成。")

    if st.session_state.final_report:
        st.markdown(st.session_state.final_report)
        st.download_button(
            "下载完整报告 Markdown",
            data=st.session_state.final_report,
            file_name="career_agent_report.md",
            mime="text/markdown",
        )
    else:
        st.info("点击“生成完整求职优化报告”后，这里会显示可下载报告。")

    st.divider()
    st.markdown("### Agent 运行日志")
    log_text = "\n".join(st.session_state.logs) if st.session_state.logs else "[System] 等待任务启动..."
    st.code(log_text, language="text")


def main() -> None:
    configure_page()
    init_state()

    st.title(APP_TITLE)
    st.caption("单文件简化版 Demo：简历诊断 + 岗位匹配 + 简历优化 + 模拟面试 + 回答评分")

    api_key, model, demo_mode = render_sidebar()
    tab1, tab2, tab3, tab4 = st.tabs(["1. 简历与 JD", "2. 简历分析", "3. 模拟面试", "4. 报告与日志"])

    with tab1:
        render_input_tab()
    with tab2:
        render_analysis_tab(api_key=api_key, model=model, demo_mode=demo_mode)
    with tab3:
        render_interview_tab(api_key=api_key, model=model, demo_mode=demo_mode)
    with tab4:
        render_report_tab(api_key=api_key, model=model, demo_mode=demo_mode)

    st.divider()
    st.caption("注意：本 Demo 可优化表达与训练面试，但不应编造学历、公司、项目、证书或量化成果。")


if __name__ == "__main__":
    main()
